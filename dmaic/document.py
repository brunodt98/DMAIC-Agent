"""
document.py — Gerador do documento Word (Project Charter DMAIC).

Responsável por montar o arquivo .docx com qualidade Black Belt,
incluindo: capa, ficha de identificação, situação consolidada, SIPOC,
Ishikawa, Plano de Campo, Síntese da Sessão e próximos passos.
"""

import io
import re
import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from dmaic.config import ETAPAS_META, ROADMAP


# ─────────────────────────────────────────────────────────────────
# UTILITÁRIOS DE FORMATAÇÃO
# ─────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str) -> None:
    """Define cor de fundo de uma célula de tabela."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


class DocBuilder:
    """Wrapper com helpers para construir o documento de forma fluente."""

    def __init__(self):
        self.doc = Document()
        for sec in self.doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2)
            sec.left_margin = sec.right_margin = Cm(2.5)

    # ── Cabeçalho colorido de seção ───────────────────────────────
    def section_header(self, text: str, hex_color: str) -> None:
        t = self.doc.add_table(rows=1, cols=1)
        t.autofit = False
        t.columns[0].width = Cm(18)
        cell = t.rows[0].cells[0]
        _set_cell_bg(cell, hex_color)
        run = cell.paragraphs[0].add_run(text)
        run.font.bold  = True
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        self.doc.add_paragraph()

    # ── Tabela chave → valor ──────────────────────────────────────
    def kv_table(self, pairs: list[tuple[str, str]],
                 label_color: str = "EBF5FB") -> None:
        t = self.doc.add_table(rows=0, cols=2)
        t.style    = "Table Grid"
        t.autofit  = False
        t.columns[0].width = Cm(5.5)
        t.columns[1].width = Cm(12.5)
        for label, value in pairs:
            row = t.add_row()
            _set_cell_bg(row.cells[0], label_color)
            rl = row.cells[0].paragraphs[0].add_run(label)
            rl.font.bold = True
            rl.font.size = Pt(8)
            display = value if value else "Não informado"
            rr = row.cells[1].paragraphs[0].add_run(display)
            rr.font.size = Pt(9)
            if not value:
                rr.font.color.rgb = RGBColor(180, 180, 180)
                rr.font.italic    = True
        self.doc.add_paragraph()

    # ── Parágrafo de texto simples ────────────────────────────────
    def text(self, content: str, size: int = 9, bold: bool = False,
             italic: bool = False, hex_color: str = None) -> None:
        p   = self.doc.add_paragraph()
        run = p.add_run(content)
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if hex_color:
            run.font.color.rgb = RGBColor(
                *[int(hex_color[i:i+2], 16) for i in (0, 2, 4)]
            )

    # ── Bullet point ─────────────────────────────────────────────
    def bullet(self, content: str, size: int = 9) -> None:
        p   = self.doc.add_paragraph(style="List Bullet")
        run = p.add_run(content)
        run.font.size = Pt(size)

    # ── Numbered list item ────────────────────────────────────────
    def numbered(self, content: str, size: int = 9) -> None:
        p   = self.doc.add_paragraph(style="List Number")
        run = p.add_run(content)
        run.font.size = Pt(size)

    # ── Serializa o documento para bytes ─────────────────────────
    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.read()


# ─────────────────────────────────────────────────────────────────
# RENDERIZAÇÃO DO BLOCO DE PLANO/SÍNTESE
# ─────────────────────────────────────────────────────────────────
def _render_chat_block(builder: DocBuilder, texto: str,
                       header_color: tuple[int, int, int]) -> None:
    """
    Renderiza um bloco de texto do chat (Plano de Campo ou Síntese)
    aplicando formatação adequada linha a linha.
    """
    for linha in texto.split("\n"):
        linha = linha.strip()
        if not linha:
            continue
        limpa = linha.replace("**", "").strip()
        limpa = re.sub(r"[📋📊]", "", limpa).strip()

        # Cabeçalho de seção dentro do bloco
        if any(k in linha.upper() for k in
               ["PLANO DE CAMPO", "SÍNTESE DA SESSÃO", "SÍNTESE"]):
            p   = builder.doc.add_paragraph()
            run = p.add_run(limpa)
            run.font.bold  = True
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(*header_color)
        elif re.match(r"^[•\-\*–]", linha):
            builder.bullet(limpa.lstrip("•-*– "))
        elif re.match(r"^\d+[\.\)]", linha):
            builder.numbered(re.sub(r"^\d+[\.\)]\s*", "", limpa))
        elif limpa:
            builder.text(limpa)


# ─────────────────────────────────────────────────────────────────
# FUNÇÃO PRINCIPAL
# ─────────────────────────────────────────────────────────────────
def gerar_word() -> bytes:
    """
    Gera e retorna o documento Word como bytes.
    Lê dados de st.session_state (projeto, meta, chat, etapa).
    """
    proj  = st.session_state.projeto
    meta  = st.session_state.meta
    chat  = st.session_state.chat
    etapa = st.session_state.etapa
    b     = DocBuilder()

    etapa_doc = ETAPAS_META.get(etapa, {}).get("doc", "DMAIC")
    agora     = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")

    # ── Capa ──────────────────────────────────────────────────────
    t_capa = b.doc.add_table(rows=1, cols=2)
    t_capa.autofit = False
    t_capa.columns[0].width = Cm(11)
    t_capa.columns[1].width = Cm(7)
    for cell in t_capa.rows[0].cells:
        _set_cell_bg(cell, "1B2A4A")

    r0 = t_capa.rows[0].cells[0].paragraphs[0].add_run(
        f"🏢 {meta.get('empresa', '—')}  |  DMAIC PROJECT CHARTER")
    r0.font.bold  = True
    r0.font.size  = Pt(12)
    r0.font.color.rgb = RGBColor(255, 255, 255)

    r1 = t_capa.rows[0].cells[1].paragraphs[0].add_run(agora)
    r1.font.size  = Pt(9)
    r1.font.color.rgb = RGBColor(180, 180, 180)
    t_capa.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    b.doc.add_paragraph()

    titulo_p = b.doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_r = titulo_p.add_run(
        proj.get("titulo", proj.get("problema", "Projeto DMAIC"))[:80])
    titulo_r.font.size  = Pt(16)
    titulo_r.font.bold  = True
    titulo_r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    sub_p = b.doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_p.add_run(f"Etapa atual: {etapa_doc}")
    sub_r.font.size   = Pt(10)
    sub_r.font.italic = True
    sub_r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    b.doc.add_paragraph()

    # ── Identificação ─────────────────────────────────────────────
    b.section_header("📋 IDENTIFICAÇÃO DO PROJETO", "1B2A4A")
    b.kv_table([
        ("Responsável",  meta.get("responsavel", "—")),
        ("Patrocinador", meta.get("patrocinador", "—")),
        ("Área",         meta.get("area", "—")),
        ("Prazo",        proj.get("prazo", "—")),
        ("Nº Projeto",   meta.get("numero", "—")),
        ("Início",       meta.get("inicio", "—")),
    ], label_color="EBF5FB")

    # ── Situação atual consolidada ────────────────────────────────
    b.section_header("🎯 SITUAÇÃO ATUAL CONSOLIDADA", "1E3A5F")
    campos_sit = [
        ("Problema identificado",  proj.get("problema", "")),
        ("Impacto",                proj.get("impacto", "")),
        ("Objetivo do projeto",    proj.get("objetivo", "")),
        ("Escopo",                 proj.get("escopo", "")),
        ("Equipe envolvida",       proj.get("equipe", "")),
        ("Situação atual (dados)", proj.get("situacao_atual", "")),
        ("Meta quantitativa",      proj.get("meta", "")),
        ("Frequência",             proj.get("frequencia", "")),
        ("Baseline",               proj.get("baseline", "")),
        ("VOC — Voz do Cliente",   proj.get("voc", "")),
        ("CTQ",                    proj.get("ctq", "")),
    ]
    b.kv_table([(l, v) for l, v in campos_sit if v], label_color="D6EAF8")

    # ── SIPOC ─────────────────────────────────────────────────────
    campos_sipoc = ["sipoc_fornecedores", "sipoc_entradas", "sipoc_processo",
                    "sipoc_saidas", "sipoc_clientes"]
    if any(proj.get(c, "") for c in campos_sipoc):
        b.section_header("🗺️ SIPOC — MAPA DO PROCESSO", "21618C")
        b.kv_table([
            ("Fornecedores (S)", proj.get("sipoc_fornecedores", "")),
            ("Entradas (I)",     proj.get("sipoc_entradas", "")),
            ("Processo (P)",     proj.get("sipoc_processo", "")),
            ("Saídas (O)",       proj.get("sipoc_saidas", "")),
            ("Clientes (C)",     proj.get("sipoc_clientes", "")),
        ], label_color="D6EAF8")

    # ── Ferramentas aplicadas ─────────────────────────────────────
    ferramentas = list(st.session_state.get("ferramentas_usadas", []))
    extras_str  = proj.get("ferramentas_aplicadas", "")
    if extras_str:
        for f in extras_str.split(","):
            f = f.strip()
            if f and f not in ferramentas:
                ferramentas.append(f)

    if ferramentas:
        b.section_header("🛠️ FERRAMENTAS APLICADAS NESTA SESSÃO", "1A5276")
        for ferr in ferramentas:
            b.bullet(ferr)
        b.doc.add_paragraph()

    # ── Ishikawa ──────────────────────────────────────────────────
    campos_ishi = ["ishikawa_metodo", "ishikawa_maquina", "ishikawa_mao_obra",
                   "ishikawa_material", "ishikawa_meio_ambiente", "ishikawa_medicao"]
    if any(proj.get(c, "") for c in campos_ishi):
        b.section_header("🐟 DIAGRAMA DE ISHIKAWA — CAUSAS LEVANTADAS", "7D6608")
        b.kv_table([
            ("Método",        proj.get("ishikawa_metodo", "")),
            ("Máquina",       proj.get("ishikawa_maquina", "")),
            ("Mão de Obra",   proj.get("ishikawa_mao_obra", "")),
            ("Material",      proj.get("ishikawa_material", "")),
            ("Meio Ambiente", proj.get("ishikawa_meio_ambiente", "")),
            ("Medição",       proj.get("ishikawa_medicao", "")),
        ], label_color="FEF9E7")

    # ── Hipóteses e causa raiz ────────────────────────────────────
    campos_analise = ["causas", "causa_raiz", "evidencias",
                      "cinco_porques", "matriz_priorizacao"]
    if any(proj.get(c, "") for c in campos_analise):
        b.section_header("🔍 ANÁLISE — HIPÓTESES E CAUSA RAIZ", "7D6608")
        b.kv_table([
            ("Possíveis causas",       proj.get("causas", "")),
            ("Hipótese de causa raiz", proj.get("causa_raiz", "")),
            ("Evidências disponíveis", proj.get("evidencias", "")),
            ("5 Porquês",              proj.get("cinco_porques", "")),
            ("Matriz de Priorização",  proj.get("matriz_priorizacao", "")),
        ], label_color="FEF9E7")

    # ── Plano de Campo ────────────────────────────────────────────
    planos = [m["content"] for m in chat
              if m["role"] == "assistant"
              and "plano de campo" in m["content"].lower()]

    if planos:
        b.section_header("📋 PLANO DE CAMPO — O QUE FAZER AGORA", "C0392B")
        b.text("Execute as tarefas abaixo antes de retornar à sessão com o consultor:",
               italic=True, hex_color="666666")
        b.doc.add_paragraph()
        _render_chat_block(b, planos[-1], (0xC0, 0x39, 0x2B))
        b.doc.add_paragraph()

    # ── Síntese da Sessão ─────────────────────────────────────────
    sinteses = [m["content"] for m in chat
                if m["role"] == "assistant"
                and "síntese da sessão" in m["content"].lower()]

    if sinteses and planos:
        b.section_header("📊 SÍNTESE DA SESSÃO", "1A5276")
        ultima = sinteses[-1]
        idx = ultima.lower().find("síntese da sessão")
        bloco = ultima[idx:] if idx >= 0 else ultima
        _render_chat_block(b, bloco, (0x1A, 0x52, 0x76))
        b.doc.add_paragraph()

    # ── Próximos passos ───────────────────────────────────────────
    b.section_header("🚀 PRÓXIMOS PASSOS — APÓS RETORNO DO CAMPO", "1E8449")
    b.text("Quando retornar com as informações levantadas, o consultor irá conduzir:",
           italic=True, hex_color="666666")
    b.doc.add_paragraph()
    for passo in ROADMAP.get(etapa, ROADMAP["definir"]):
        b.bullet(passo)
    b.doc.add_paragraph()

    # ── Melhorias e Controle ──────────────────────────────────────
    campos_melhoria = ["solucoes", "solucao_escolhida", "fmea",
                       "plano_acao_5w2h", "resultados"]
    if any(proj.get(c, "") for c in campos_melhoria):
        b.section_header("⚡ MELHORIAS E CONTROLE", "6C3483")
        b.kv_table([
            ("Soluções propostas",    proj.get("solucoes", "")),
            ("Solução escolhida",     proj.get("solucao_escolhida", "")),
            ("FMEA",                  proj.get("fmea", "")),
            ("Plano de Ação (5W2H)",  proj.get("plano_acao_5w2h", "")),
            ("Recursos necessários",  proj.get("recursos", "")),
            ("Riscos identificados",  proj.get("riscos", "")),
            ("Resultados esperados",  proj.get("resultados", "")),
            ("KPIs de monitoramento", proj.get("indicadores", "")),
            ("Padronização (SOP)",    proj.get("padronizacao", "")),
            ("Lições aprendidas",     proj.get("licoes", "")),
        ], label_color="F4ECF7")

    # ── Rodapé ────────────────────────────────────────────────────
    rod = b.doc.add_paragraph(
        f"DMAIC Agent  ·  {meta.get('empresa', '—')}  ·  "
        f"Responsável: {meta.get('responsavel', '—')}  ·  "
        f"Gerado em {agora}")
    rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rod.runs[0].font.size      = Pt(7)
    rod.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    return b.to_bytes()
