"""
ui/sidebar.py — Barra lateral: API key, modelo, upload, progresso e exportação.
"""

import streamlit as st

from dmaic.ai import extrair_dados, detectar_etapa
from dmaic.document import gerar_word
from dmaic.state import reset_state
from dmaic.config import GROQ_MODELS, ETAPAS, ETAPAS_META


def _ler_docx_bytes(file_bytes: bytes) -> dict:
    """Lê tabelas e texto de um .docx existente."""
    import io
    from docx import Document

    doc   = Document(io.BytesIO(file_bytes))
    dados = {}
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                k = row.cells[0].text.strip()
                v = row.cells[1].text.strip()
                if k and v and len(v) > 2:
                    dados[k[:60]] = v
    texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return {"tabelas": dados, "texto": texto[:3000]}


def render_sidebar() -> None:
    with st.sidebar:
        st.markdown("## 🎯 DMAIC Agent")
        st.caption("Consultor Six Sigma com IA")
        st.divider()

        # ── API Key e modelo ──────────────────────────────────────
        key_in = st.text_input(
            "🔑 API Key Groq", type="password",
            help="Gratuita em https://console.groq.com/keys")
        if key_in:
            st.session_state.groq_key = key_in
            st.success("✅ Conectado!")

        st.session_state.model = st.selectbox(
            "Modelo LLM", GROQ_MODELS, label_visibility="collapsed")

        st.divider()

        # ── Retomar projeto existente ─────────────────────────────
        st.markdown("**📂 Retomar projeto**")
        uploaded = st.file_uploader(
            "Carregar .docx gerado anteriormente",
            type=["docx"], label_visibility="collapsed")

        if uploaded and st.button("📥 Carregar e continuar",
                                  use_container_width=True):
            import json
            from dmaic.ai import ai_chat

            with st.spinner("Lendo documento..."):
                lido = _ler_docx_bytes(uploaded.read())

            ctx = (
                f"O usuário retornou com um documento DMAIC existente.\n"
                f"Dados encontrados:\n"
                f"{json.dumps(lido['tabelas'], ensure_ascii=False)[:2000]}\n"
                f"Texto adicional: {lido['texto'][:600]}\n\n"
                f"Analise o que já foi feito, identifique onde o projeto está "
                f"e o que está incompleto. Faça um briefing em 3 parágrafos: "
                f"(1) o que já foi construído, (2) o que ficou pendente, "
                f"(3) próximo passo mais importante. "
                f"Depois conduza diretamente. NÃO faça mais de 1 pergunta."
            )
            for k, v in lido["tabelas"].items():
                st.session_state.projeto[k[:40]] = v

            with st.spinner("Analisando projeto retomado..."):
                resposta = ai_chat(extra_ctx=ctx)

            st.session_state.chat             = [{"role": "assistant", "content": resposta}]
            st.session_state.pronto           = True
            st.session_state.aguardando_campo = False
            st.rerun()

        st.divider()

        # ── Status e progresso (só após onboarding) ───────────────
        if st.session_state.pronto:
            _render_status()
            st.divider()
            _render_export()

        st.divider()

        if st.button("🔄 Novo projeto", use_container_width=True):
            reset_state()
            st.rerun()

        st.caption("DMAIC Agent v4 · FATEC Cotia × Outtech Services IT")


def _render_status() -> None:
    """Exibe status atual e progresso das etapas."""
    if st.session_state.get("aguardando_campo"):
        st.warning("⏸️ **Aguardando retorno do campo**")
        st.caption("Carregue o .docx preenchido ou continue a conversa com os dados.")
    else:
        st.info(f"▶️ Etapa: **{st.session_state.etapa.upper()}**")

    st.markdown("**Progresso:**")
    atual   = st.session_state.etapa
    atual_i = ETAPAS.index(atual)
    for i, e in enumerate(ETAPAS):
        emoji = ETAPAS_META[e]["emoji"]
        label = ETAPAS_META[e]["label"]
        if i < atual_i:
            st.markdown(f"✅ {emoji} {label}")
        elif e == atual:
            st.markdown(f"▶️ **{emoji} {label.upper()}**")
        else:
            st.markdown(f"⏳ {emoji} {label}")

    ferramentas = st.session_state.get("ferramentas_usadas", [])
    if ferramentas:
        st.divider()
        st.markdown("**🛠️ Ferramentas aplicadas:**")
        for f in ferramentas:
            st.caption(f"• {f}")


def _render_export() -> None:
    """Botões de atualização e exportação do documento Word."""
    st.markdown("**📥 Exportar**")

    if st.button("🔄 Atualizar dados", use_container_width=True):
        with st.spinner("Extraindo dados da conversa..."):
            dados = extrair_dados(st.session_state.chat)
            if dados:
                st.session_state.projeto.update(dados)
            st.session_state.etapa = detectar_etapa(st.session_state.chat)
        st.success(f"✅ {len(dados)} campos atualizados!")

    if st.button("📄 Gerar Word", use_container_width=True, type="primary"):
        with st.spinner("Gerando documento Black Belt..."):
            dados = extrair_dados(st.session_state.chat)
            if dados:
                st.session_state.projeto.update(dados)
            st.session_state.word_bytes = gerar_word()
        st.success("✅ Documento pronto!")

    if st.session_state.word_bytes:
        nome = st.session_state.projeto.get(
            "titulo",
            st.session_state.projeto.get("problema", "projeto")
        )[:25].replace(" ", "_")

        tem_plano = any(
            "plano de campo" in m["content"].lower()
            for m in st.session_state.chat
            if m["role"] == "assistant"
        )
        if tem_plano:
            st.info("📋 Documento de campo pronto — leve para o campo!")

        st.download_button(
            "⬇️ Baixar .docx",
            data=st.session_state.word_bytes,
            file_name=f"DMAIC_{nome}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary" if tem_plano else "secondary",
        )
